from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

def generate_word_doc(policy_name: str, content: str, version: str = "1.0", 
                     owner: str = "CISO", classification: str = "Internal", 
                     output_dir: str = "generated_docs") -> str:
    """
    Generates a professional Word document with Metadata Table and formatted content.
    Returns the absolute path to the generated file.
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    filename = f"{policy_name.replace(' ', '_')}_v{version}.docx"
    filepath = os.path.abspath(os.path.join(output_dir, filename))
    
    doc = Document()
    
    # 1. Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "AssuRisk Compliance Hub"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 2. Title Page
    title = doc.add_heading(policy_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Spacer
    
    # 3. Metadata Table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid' # 'Light Grid Accent 1' might not exist in default template, using standard
    
    metadata = [
        ('Document ID:', f"POL-{policy_name.replace(' ', '-').upper()[:10]}-001"),
        ('Version:', version),
        ('Effective Date:', datetime.now().strftime('%Y-%m-%d')),
        ('Owner:', owner),
        ('Classification:', classification),
        ('Next Review:', '2027-01-17') # Placeholder logic
    ]
    
    for i, (key, value) in enumerate(metadata):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value
    
    doc.add_page_break()
    
    # 4. Parse Content (Markdown-ish to Docx)
    # We treat content as a blob and try to format headers
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            p = doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)
            
    doc.save(filepath)
    return filepath
