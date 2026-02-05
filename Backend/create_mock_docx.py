from docx import Document

doc = Document()
doc.add_heading('Access Control Policy (Gold Standard)', 0)
doc.add_paragraph('Control ID: A.5.15')
doc.add_paragraph('This is a master template. It is immutable.')
doc.add_paragraph('Company Name: {{company_name}}')
doc.add_paragraph('MFA Tool: {{mfa_tool}}')

doc.save(r"C:\Users\dhank\OneDrive\Documents\Compliance_Product\Backend\src\assets\templates\master\ISO27001_A.5.15_Access_Control.docx")
print("Mock Template Created")
